"""
文档加载器模块
支持加载 txt、md、csv、pdf、docx 格式的文档，统一输出为 Document 列表

安全措施：
- 文件大小限制（10MB）
- 扩展名白名单（仅 .txt/.md/.csv/.pdf/.docx）
- 文件名净化（防止路径遍历）
- MIME 类型验证（读取文件头）
"""
import logging
import re
from pathlib import Path

from langchain_core.documents import Document

logger = logging.getLogger("ai_rd_agent")


# 支持的文件格式及其对应的加载器
SUPPORTED_EXTENSIONS = {
    ".txt": "text",
    ".md": "text",
    ".csv": "text",
    ".pdf": "pdf",
    ".docx": "docx",
}

# MIME 类型签名（文件头 magic bytes）
_MIME_SIGNATURES: dict[str, bytes] = {
    ".pdf": b"%PDF",
    ".docx": b"PK",  # docx 是 ZIP 格式
}


def sanitize_filename(filename: str) -> str:
    """净化文件名 — 移除路径遍历和特殊字符

    保留中文字符、字母、数字、连字符、下划线、点号，其余替换为下划线。
    """
    # 仅保留文件名部分（移除目录路径）
    name = Path(filename).name
    # 替换非安全字符
    safe = re.sub(r'[^\w.一-鿿-]', '_', name)
    return safe if safe else "unnamed_file"


class DocumentLoader:
    """统一文档加载器 — 根据文件扩展名自动选择解析方式

    使用示例：
        loader = DocumentLoader()
        docs = loader.load("data/docs/test_report.txt")
        for batch in loader.load_directory_batch("data/docs/", batch_size=5):
            print(f"处理了 {len(batch)} 个文档")
        docs = loader.load_directory("data/docs/")
    """

    # 允许加载的最大文件大小（10MB），防止内存溢出
    MAX_FILE_SIZE = 10 * 1024 * 1024

    # 截断保护上限（50MB）：文本类文件超过 MAX_FILE_SIZE 但不超过该值时，
    # 截断到 MAX_FILE_SIZE 索引并在 metadata 标记 truncated（降级可用）；
    # 超过该值或二进制格式（pdf/docx 无法部分解析）超限时直接拒绝。
    TRUNCATE_THRESHOLD = 50 * 1024 * 1024

    # 文本类扩展名（可安全按字节截断）
    _TEXT_EXTS = frozenset({".txt", ".md", ".csv"})

    def load(self, file_path: str) -> list[Document]:
        """加载单个文档文件

        Args:
            file_path: 文档文件的绝对路径

        Returns:
            LangChain Document 列表（每个 Document 包含 page_content 和 metadata）
        """
        path = Path(file_path)

        # 检查文件是否存在
        if not path.exists():
            raise FileNotFoundError(f"文档不存在: {file_path}")

        # 检查文件大小
        file_size = path.stat().st_size
        if file_size > self.MAX_FILE_SIZE:
            is_text = path.suffix.lower() in self._TEXT_EXTS
            if is_text and file_size <= self.TRUNCATE_THRESHOLD:
                # 文本类降级：截断索引（truncated 标记在 _load_textlike 里写入 metadata）
                logger.warning(
                    f"文件 {path.name} 大小 {self._format_size(file_size)} 超过 "
                    f"{self._format_size(self.MAX_FILE_SIZE)}，截断索引前 "
                    f"{self._format_size(self.MAX_FILE_SIZE)} 内容"
                )
            else:
                raise ValueError(
                    f"文件过大: {self._format_size(file_size)}，"
                    f"最大允许 {self._format_size(self.MAX_FILE_SIZE)}"
                )

        # 根据扩展名选择加载方式
        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            supported = ", ".join(SUPPORTED_EXTENSIONS.keys())
            raise ValueError(f"不支持的文件格式 '{ext}'。支持的格式: {supported}")

        # MIME 类型验证（支持的类型才读取文件头）
        if ext in _MIME_SIGNATURES:
            expected_header = _MIME_SIGNATURES[ext]
            with open(path, "rb") as f:
                actual_header = f.read(len(expected_header))
            if not actual_header.startswith(expected_header):
                raise ValueError(
                    f"文件 '{path.name}' 的扩展名与内容不匹配。"
                    f"期望 {ext} 格式，但文件头为 {actual_header[:8].hex()}"
                )

        logger.info(f"加载文档: {path.name} ({self._format_size(file_size)})")

        if ext == ".txt":
            return self._load_textlike(path, "txt")
        elif ext == ".md":
            return self._load_textlike(path, "md")
        elif ext == ".csv":
            return self._load_textlike(path, "csv")
        elif ext == ".pdf":
            return self._load_pdf(path)
        elif ext == ".docx":
            return self._load_docx(path)
        else:
            return []

    def load_directory(self, dir_path: str) -> list[Document]:
        """加载目录下所有支持的文档

        Args:
            dir_path: 文档目录路径

        Returns:
            所有文档的 Document 列表
        """
        all_docs = []
        dir_path = Path(dir_path)

        if not dir_path.exists():
            logger.warning(f"文档目录不存在: {dir_path}")
            return all_docs

        for ext in SUPPORTED_EXTENSIONS:
            for file_path in dir_path.glob(f"*{ext}"):
                try:
                    docs = self.load(str(file_path))
                    all_docs.extend(docs)
                except Exception as e:
                    logger.error(f"加载失败 {file_path.name}: {e}")

        logger.info(f"目录加载完成: {len(all_docs)} 个文档片段")
        return all_docs

    def load_directory_batch(self, dir_path: str, batch_size: int = 10) -> list[list[Document]]:
        """分批加载目录下的文档（适用于流式处理）

        每次产出 batch_size 个文件的文档片段列表，避免大目录下全量加载到内存。
        配合 RAGPipeline.ingest_directory_batch 使用。

        Args:
            dir_path: 文档目录路径
            batch_size: 每批文件数（默认 10）

        Yields:
            每批文档的 Document 列表
        """
        dir_path = Path(dir_path)

        if not dir_path.exists():
            logger.warning(f"文档目录不存在: {dir_path}")
            return

        # 收集所有支持的文件
        all_files: list[Path] = []
        for ext in SUPPORTED_EXTENSIONS:
            all_files.extend(sorted(dir_path.glob(f"*{ext}")))

        total_files = len(all_files)
        if total_files == 0:
            logger.info(f"目录中没有支持的文档: {dir_path}")
            return

        logger.info(f"开始分批加载目录: {total_files} 个文件 (批次大小={batch_size})")

        for start in range(0, total_files, batch_size):
            batch_files = all_files[start:start + batch_size]
            batch_docs: list[Document] = []
            for file_path in batch_files:
                try:
                    docs = self.load(str(file_path))
                    batch_docs.extend(docs)
                except Exception as e:
                    logger.error(f"加载失败 {file_path.name}: {e}")

            batch_num = start // batch_size + 1
            total_batches = (total_files + batch_size - 1) // batch_size
            logger.info(f"  批次 {batch_num}/{total_batches}: {len(batch_files)} 个文件 → {len(batch_docs)} 个文档片段")
            yield batch_docs

    # ---- 各格式加载器 ----

    # charset_normalizer 检测结果白名单：这些编码可信直接采用。
    # cp949（韩文）等与 GB 系字节结构高度歧义，短文本下常误报，不在白名单。
    _TRUSTED_DETECTED_ENCODINGS = frozenset({
        "utf_8", "utf-16", "utf-16_be", "utf-16_le", "utf_16", "big5", "ascii",
    })

    def _decode_text(self, raw: bytes, filename: str) -> tuple[str, str]:
        """解码文本字节：utf-8 → gb18030（中文主场景）→ 检测器白名单 → 容错兜底

        Returns:
            (解码文本, 实际编码名)
        """
        try:
            return raw.decode("utf-8"), "utf-8"
        except UnicodeDecodeError:
            pass

        # gb18030 是 GB2312/GBK 超集，中文文本通常可严格解码（无替换字符即为强信号）
        try:
            text = raw.decode("gb18030")
            if "\ufffd" not in text:
                return text, "gb18030"
        except UnicodeDecodeError:
            pass

        # 检测回退（charset_normalizer 为 httpx 传递依赖，无新增安装），
        # 结果需在白名单内（排除 cp949 等歧义误报）
        from charset_normalizer import from_bytes
        match = from_bytes(raw).best()
        if match is not None and match.encoding:
            norm = match.encoding.lower().replace("-", "_")
            if norm in self._TRUSTED_DETECTED_ENCODINGS:
                logger.info(f"文件 {filename} 非 UTF-8，检测到编码 {match.encoding}，自动回退解码")
                return str(match), match.encoding

        # 终极兜底：gb18030 + replace，保证可读不中断
        logger.warning(f"文件 {filename} 编码检测失败，按 gb18030 容错解码")
        return raw.decode("gb18030", errors="replace"), "gb18030(replace)"

    def _load_textlike(self, path: Path, file_type: str) -> list[Document]:
        """加载文本类文件（.txt / .md / .csv）

        - 编码：utf-8 → charset_normalizer 检测 → gb18030 容错三级回退
        - 超大文件：超过 MAX_FILE_SIZE 的文本按字节截断到上限，metadata 标记 truncated
        """
        file_size = path.stat().st_size
        truncated = file_size > self.MAX_FILE_SIZE

        with open(path, "rb") as f:
            raw = f.read(self.MAX_FILE_SIZE if truncated else -1)

        content, encoding = self._decode_text(raw, path.name)

        # CSV 文件特殊处理：前 100 行预览 + 结构化数据标记
        if file_type == "csv":
            import csv
            import io
            reader = csv.reader(io.StringIO(content))
            rows = list(reader)
            header = rows[0] if rows else []
            preview_rows = rows[:100]
            content = (
                f"[CSV 数据] {path.name}\n"
                f"列: {', '.join(header)}\n"
                f"行数: {len(rows)}\n"
                f"--- 前 {len(preview_rows)} 行 ---\n"
                + "\n".join(",".join(row) for row in preview_rows)
            )

        metadata = {
            "source": str(path),
            "filename": path.name,
            "file_type": file_type,
            "encoding": encoding,
        }
        if truncated:
            metadata["truncated"] = True
            content += (
                f"\n\n[注意: 文件超过 {self._format_size(self.MAX_FILE_SIZE)}，"
                f"仅索引前 {self._format_size(self.MAX_FILE_SIZE)} 内容]"
            )

        return [Document(page_content=content, metadata=metadata)]

    def _load_pdf(self, path: Path) -> list[Document]:
        """加载 PDF 文件 — 逐页提取文本"""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("请安装 pypdf: pip install pypdf") from None

        reader = PdfReader(str(path))
        docs = []

        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                docs.append(Document(
                    page_content=text.strip(),
                    metadata={
                        "source": str(path),
                        "filename": path.name,
                        "file_type": "pdf",
                        "page": i + 1,
                    }
                ))

        logger.info(f"  PDF 解析完成: {len(docs)} 页")
        return docs

    def _load_docx(self, path: Path) -> list[Document]:
        """加载 Word 文档"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx") from None

        doc = DocxDocument(str(path))

        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 也提取表格中的文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text for cell in row.cells if cell.text.strip()
                )
                if row_text.strip():
                    paragraphs.append(row_text)

        full_text = "\n".join(paragraphs)

        return [Document(
            page_content=full_text,
            metadata={
                "source": str(path),
                "filename": path.name,
                "file_type": "docx",
            }
        )]

    @staticmethod
    def _format_size(size_bytes: int) -> str:
        """将字节数格式化为人类可读的大小"""
        for unit in ["B", "KB", "MB"]:
            if size_bytes < 1024:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f} GB"
